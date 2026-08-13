"""Attachment text extraction for A1 and B3.

The flagship A1 (bank-detail change) was blind to attachments — but the changed
IBAN almost always lives inside an attached PDF or Word invoice, or an image. This
module pulls text out of those so A1 sees it, and decodes QR codes so B3 (quishing)
can run the URL checks on what the code points at.

Everything degrades gracefully: a missing system library (tesseract for OCR, zbar
for QR) disables just that path and never raises. All work is bounded — we cap
bytes and page counts so a malicious 500-page PDF can't exhaust a worker.
"""

from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger("envelock.attachments")

#: Bound the work a single attachment can cause.
_MAX_BYTES = 25 * 1024 * 1024
_MAX_PDF_PAGES = 30
_URL_RE = re.compile(r"https?://[^\s\"'<>)]+", re.IGNORECASE)

_PDF_MIMES = {"application/pdf"}
_DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_IMAGE_PREFIX = "image/"


def _pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages[:_MAX_PDF_PAGES]:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        logger.debug("pdf extract failed: %s", exc)
        return ""


def _docx_text(data: bytes) -> str:
    try:
        import docx  # python-docx

        document = docx.Document(io.BytesIO(data))
        lines = [p.text for p in document.paragraphs]
        # Bank details are frequently in a table cell, not a paragraph.
        for table in document.tables:
            for row in table.rows:
                lines.extend(cell.text for cell in row.cells)
        return "\n".join(t for t in lines if t)
    except Exception as exc:  # noqa: BLE001
        logger.debug("docx extract failed: %s", exc)
        return ""


def _image_ocr(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(data))
        return pytesseract.image_to_string(image) or ""
    except Exception as exc:  # noqa: BLE001 — tesseract binary may be absent
        logger.debug("ocr failed (is tesseract installed?): %s", exc)
        return ""


def _image_qr_urls(data: bytes) -> tuple[str, ...]:
    try:
        from PIL import Image
        from pyzbar.pyzbar import decode

        image = Image.open(io.BytesIO(data))
        urls: list[str] = []
        for code in decode(image):
            try:
                text = code.data.decode("utf-8", "ignore")
            except Exception:  # noqa: BLE001, S112
                continue
            urls.extend(_URL_RE.findall(text))
        return tuple(dict.fromkeys(urls))
    except Exception as exc:  # noqa: BLE001 — zbar shared lib may be absent
        logger.debug("qr decode failed (is zbar installed?): %s", exc)
        return ()


def extract(data: bytes, *, mime: str | None, filename: str) -> tuple[str, tuple[str, ...]]:
    """Return (extracted_text, qr_urls) for one attachment. Best-effort and bounded."""
    if not data or len(data) > _MAX_BYTES:
        return "", ()
    mime = (mime or "").lower()
    name = filename.lower()

    if mime in _PDF_MIMES or name.endswith(".pdf"):
        return _pdf_text(data), ()
    if mime in _DOCX_MIMES or name.endswith(".docx"):
        return _docx_text(data), ()
    if mime.startswith(_IMAGE_PREFIX) or name.endswith(
        (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp")
    ):
        return _image_ocr(data), _image_qr_urls(data)
    return "", ()


__all__ = ["extract"]
